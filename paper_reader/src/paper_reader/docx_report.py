"""Generate concise DOCX reports with rendered LaTeX math and embedded figures."""

from __future__ import annotations

import re
from io import BytesIO
from lxml import etree
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from paper_reader.reporting import Discovery, _slugify


_M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
_FIGURE_CAPTION_RE = re.compile(r"Figure\s+(\d+)[:\.]", re.IGNORECASE)

# Minimum image area for a standalone figure (skip sub-plots, icons, glyphs).
# Most real figures are at least 800px wide. Set high to prefer page renders for composites.
_MIN_IMAGE_AREA = 800 * 400


def _latex_to_omml(latex: str) -> etree._Element | None:
    """Convert a LaTeX string to an Office Math ML (OMML) element for Word."""
    try:
        import latex2mathml.converter
        from mathml2omml import convert as mathml2omml_convert

        mathml_str = latex2mathml.converter.convert(latex)
        omml_str = mathml2omml_convert(mathml_str)
        wrapped = f'<root xmlns:m="{_M_NS}">{omml_str}</root>'
        root = etree.fromstring(wrapped.encode())
        return root[0]
    except Exception:
        return None


def _render_page_region(fitz_doc, page_num: int, scale: float = 2.0) -> bytes | None:
    """Render a PDF page as a PNG image."""
    try:
        import fitz
        page = fitz_doc[page_num]
        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale))
        return pix.tobytes("png")
    except Exception:
        return None


def _extract_figure_images(
    pdf_bytes: bytes, figures: list[dict], max_pages: int = 18,
) -> dict[str, bytes]:
    """Extract figure images from PDF, keyed by figure ID (e.g. 'Figure 1').

    Uses page_number from Claude's analysis when available, otherwise falls back
    to text-based caption search.  Always renders the full page as PNG — this
    handles vector figures and avoids picking up logos/sub-plot fragments.
    """
    try:
        import fitz
    except ImportError:
        return {}

    if not pdf_bytes:
        return {}

    try:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    except Exception:
        return {}

    # Build caption-page lookup as fallback
    figure_caption_pages: dict[int, int] = {}
    for page_num in range(min(len(doc), max_pages)):
        text = doc[page_num].get_text()
        for fig_num_str in _FIGURE_CAPTION_RE.findall(text):
            fig_num = int(fig_num_str)
            if fig_num not in figure_caption_pages:
                figure_caption_pages[fig_num] = page_num

    results: dict[str, bytes] = {}

    for fig in figures:
        fig_id = fig.get("figure_id", "")
        if not fig_id:
            continue

        # Determine which page to render
        page_num: int | None = None

        # Prefer page_number from Claude's vision analysis
        if "page_number" in fig:
            try:
                page_num = int(fig["page_number"])
            except (ValueError, TypeError):
                pass

        # Fallback: find by caption text
        if page_num is None:
            match = re.search(r"(\d+)", fig_id)
            if match:
                page_num = figure_caption_pages.get(int(match.group(1)))

        if page_num is None or page_num < 0 or page_num >= len(doc):
            continue

        img = _render_page_region(doc, page_num)
        if img:
            results[fig_id] = img

    doc.close()
    return results


# ── Helpers ──────────────────────────────────────────────────────────────────

def _add_formula_inline(doc: Document, formula: dict) -> None:
    """Add a compact formula block: name + rendered math + short explanation."""
    name = formula.get("name", "")
    latex = formula.get("latex", "")
    explanation = formula.get("explanation", "")

    # Name line (bold)
    if name:
        p = doc.add_paragraph()
        run = p.add_run(name)
        run.bold = True
        run.font.size = Pt(10)

    # Rendered math
    if latex:
        omath_el = _latex_to_omml(latex)
        if omath_el is not None:
            fp = doc.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp._element.append(omath_el)
        else:
            fp = doc.add_paragraph()
            fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = fp.add_run(latex.strip())
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x8C)

    # Explanation (smaller)
    if explanation:
        p = doc.add_paragraph()
        p.add_run(explanation).font.size = Pt(9)


def _add_figure_block(doc: Document, fig: dict, figure_images: dict[str, bytes]) -> None:
    """Add a compact figure block: image + caption."""
    fig_id = fig.get("figure_id", "Figure")
    sig = fig.get("significance", "")

    image_bytes = figure_images.get(fig_id)
    if image_bytes:
        img_para = doc.add_paragraph()
        img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = img_para.add_run()
        run.add_picture(BytesIO(image_bytes), width=Inches(5.0))

    if sig:
        p = doc.add_paragraph()
        run = p.add_run(f"{fig_id}: ")
        run.bold = True
        run.font.size = Pt(9)
        run2 = p.add_run(sig)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)


# ── Main writer ──────────────────────────────────────────────────────────────

def write_paper_docx(details_dir: Path, discovery: Discovery) -> Path:
    """Write a concise ~1-page DOCX report for a paper.

    Layout:
      Title + authors + link
      Research Questions  (bullets)
      Core Idea           (short intuitive explanation)
      Key Formula         (rendered OMML + explanation)
      Key Figure          (embedded image + caption)
      Limitations         (bullets)
    """
    details_dir.mkdir(parents=True, exist_ok=True)
    # Use paper title for filename (more descriptive than arxiv ID)
    title_slug = _slugify(discovery.metadata.title) if discovery.metadata.title else _slugify(discovery.paper.key)
    # Truncate to avoid filesystem limits
    if len(title_slug) > 80:
        title_slug = title_slug[:80].rstrip("-")
    docx_path = details_dir / f"{title_slug}.docx"

    doc = Document()

    # Compact default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(1)

    # Extract figure images using Claude's page_number hints
    figure_images = _extract_figure_images(
        discovery.metadata.pdf_bytes, discovery.analysis.key_figures,
    )

    # ── Title ────────────────────────────────────────────────────────────────
    title_p = doc.add_heading(discovery.metadata.title, level=1)
    title_p.runs[0].font.size = Pt(14)

    # Authors (compact)
    meta_parts = []
    if discovery.metadata.authors:
        authors_str = ", ".join(discovery.metadata.authors[:6])
        if len(discovery.metadata.authors) > 6:
            authors_str += " et al."
        meta_parts.append(authors_str)
    if discovery.metadata.venue:
        meta_parts.append(discovery.metadata.venue)

    paper_url = discovery.metadata.canonical_url or discovery.paper.canonical_url
    if meta_parts:
        p = doc.add_paragraph(" | ".join(meta_parts))
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Paper link
    if paper_url:
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        p = doc.add_paragraph()
        run = p.add_run(paper_url)
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.underline = True

    # ── Research Questions ────────────────────────────────────────────────────
    doc.add_heading("Research Questions", level=2)
    for q in discovery.analysis.research_questions:
        doc.add_paragraph(q, style="List Bullet")

    # ── Core Idea ─────────────────────────────────────────────────────────────
    if discovery.analysis.intuitive_explanation:
        doc.add_heading("Core Idea", level=2)
        # Keep it tight — at most 2 paragraphs
        paras = [p.strip() for p in discovery.analysis.intuitive_explanation.split("\n\n") if p.strip()]
        for para_text in paras[:2]:
            doc.add_paragraph(para_text)

    # ── Key Formula (only the most important 1-2) ────────────────────────────
    if discovery.analysis.key_formulas:
        doc.add_heading("Key Formula", level=2)
        for formula in discovery.analysis.key_formulas[:2]:
            _add_formula_inline(doc, formula)

    # ── Key Figure (only the most important 1) ───────────────────────────────
    if discovery.analysis.key_figures:
        doc.add_heading("Key Figure", level=2)
        _add_figure_block(doc, discovery.analysis.key_figures[0], figure_images)

    # ── Limitations ───────────────────────────────────────────────────────────
    if discovery.analysis.solution_limitations:
        doc.add_heading("Limitations", level=2)
        for item in discovery.analysis.solution_limitations[:4]:
            doc.add_paragraph(item, style="List Bullet")

    doc.save(str(docx_path))
    return docx_path
